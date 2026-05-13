"""
Generate Chapter 7 + Appendices + References + Index Word document.
Run: cd backend && uv run python ../generate_report_ch7.py
"""
import sys
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── low-level helpers ──────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=12, bold=False, italic=False,
             color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_fmt(p, size=12, line_spacing=1.5, space_before=0, space_after=8,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.line_spacing       = Pt(size * line_spacing)
    pf.space_before       = Pt(space_before)
    pf.space_after        = Pt(space_after)
    pf.alignment          = align
    for run in p.runs:
        if run.font.name in (None, ""):
            run.font.name = "Calibri"
        if run.font.size in (None, Pt(0)):
            run.font.size = Pt(size)


def add_body(doc, text, size=12, font="Calibri", line_mult=1.5,
             space_after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, font, size)
    para_fmt(p, size, line_mult, space_after=space_after, align=align)
    return p


def add_heading_custom(doc, number, title, pt_size, bold=True,
                       font="Calibri", align=WD_ALIGN_PARAGRAPH.LEFT,
                       color=(0x1e, 0x3a, 0x5f), space_before=14):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}   {title}" if number else title)
    set_font(run, font, pt_size, bold=bold, color=color)
    pf = p.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(4)
    pf.line_spacing = Pt(pt_size * 1.2)
    return p


def add_appendix_main_heading(doc, letter, title):
    """Appendix Heading 1 style: 20pt, Calibri, Bold, Left."""
    p = doc.add_paragraph()
    run = p.add_run(f"Appendix {letter}: {title}")
    set_font(run, "Calibri", 20, bold=True, color=(0x1b, 0x1b, 0x1b))
    pf = p.paragraph_format
    pf.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(18)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(24)
    return p


def add_ref_heading(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_font(run, "Calibri", 20, bold=True, color=(0x1e, 0x3a, 0x5f))
    pf = p.paragraph_format
    pf.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(18)
    pf.space_after  = Pt(8)
    pf.line_spacing = Pt(24)
    return p


def add_reference(doc, tag, text):
    p = doc.add_paragraph()
    run_tag = p.add_run(f"[{tag}]\t")
    set_font(run_tag, "Times New Roman", 12, bold=True)
    run_text = p.add_run(text)
    set_font(run_text, "Times New Roman", 12)
    pf = p.paragraph_format
    pf.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.left_indent        = Inches(0.4)
    pf.first_line_indent  = Inches(-0.4)
    pf.line_spacing       = Pt(18)
    pf.space_after        = Pt(4)
    return p


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  CONTENT
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # ══════════════════════════════════════════════════════════════════════════
    #  CHAPTER 7
    # ══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("Chapter 7: Conclusion")
    set_font(run, "Calibri", 24, bold=True, color=(0x1e, 0x3a, 0x5f))
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_body(doc,
        "This chapter presents the concluding synthesis of the AI News Desk project, "
        "drawing together the key outcomes, critical observations, and forward-looking "
        "recommendations that emerge from the development and evaluation phases. "
        "Section 7.1 provides a comprehensive project summary that recaps the platform's "
        "purpose, architecture, and feature set. Section 7.1 (Achievements and Improvements) "
        "documents the notable technical accomplishments realised during implementation, "
        "including multi-agent concurrency, extensible architecture, and auxiliary output "
        "modalities. Section 7.1 (Critical Review) offers an honest evaluation of the gaps "
        "between the project's stated requirements and its delivered implementation. "
        "Section 7.2 captures the lessons learnt across backend, frontend, and deployment "
        "dimensions. Section 7.3 identifies the most impactful future enhancements that would "
        "elevate the platform from a demonstration-grade system to a production-grade news "
        "intelligence service."
    )

    # ── 7.1  Project Summary ──────────────────────────────────────────────────
    add_heading_custom(doc, "7.1", "Project Summary", 16)

    add_body(doc,
        "The AI News Desk is a full-stack, multi-agent web application designed to aggregate, "
        "summarise, and deliver AI-generated news content across a broad range of topic "
        "categories. Built using a FastAPI Python backend and a Next.js TypeScript frontend, "
        "the system demonstrates the practical application of the OpenAI Agents SDK in "
        "constructing a production-grade, multi-agent AI pipeline accessible through a modern "
        "browser-based interface."
    )
    add_body(doc,
        "The platform was conceived to address a common friction point for news consumers: "
        "the need to manually aggregate information from multiple disparate news sources. "
        "By deploying six specialised AI agents — SEO, YouTube, Forbes, Web Search, Research, "
        "and Breaking News — the system provides a unified interface through which users can "
        "interrogate multiple information perspectives simultaneously using natural language "
        "queries. The Live News feature extends this capability by continuously monitoring and "
        "presenting auto-refreshing news across eight configurable topic categories, with a "
        "60-second refresh cycle."
    )
    add_body(doc,
        "Beyond text-based news aggregation, the system delivers a suite of auxiliary output "
        "modalities: professionally formatted downloadable PDF news reports generated by "
        "ReportLab, MP3 audio briefings produced using Google Text-to-Speech in both English "
        "and Urdu, and trend visualisation graphs rendered by Matplotlib. The Ultimate AI News "
        "Agent orchestrates the full pipeline — collecting, summarising, researching, and "
        "delivering content across all modalities in a single user-initiated run. Content safety "
        "is enforced at both input and output boundaries through a GuardrailChecker component "
        "that screens for harmful content before and after each agent execution."
    )
    add_body(doc,
        "The application is deployed on a cloud-native infrastructure: the Next.js frontend on "
        "Vercel's global CDN and the FastAPI backend on Railway's containerised cloud platform. "
        "The deployment requires no proprietary infrastructure, operates within the free-tier "
        "limits of both platforms, and is accessible from any modern web browser without "
        "client-side installation."
    )

    # ── 7.1  Achievements and Improvements ────────────────────────────────────
    add_heading_custom(doc, "7.1", "Achievements and Improvements", 16)

    add_body(doc,
        "The AI News Desk project delivered several notable achievements beyond the initial "
        "specification baseline. The following subsections describe the key technical "
        "accomplishments and qualitative improvements realised during the development phase."
    )
    add_body(doc,
        "The project successfully demonstrates a working implementation of the OpenAI Agents "
        "SDK in a production-grade application context. The six specialised agents operate "
        "concurrently using Python's asyncio.gather() mechanism, reducing multi-agent query "
        "response times compared to sequential execution. The architecture is genuinely "
        "extensible: adding a new agent type requires only the creation of a new Agent subclass "
        "and registration in the Runner's agent registry, without modification to any existing "
        "agent or API layer code."
    )
    add_body(doc,
        "The bilingual content delivery feature represents a meaningful localisation achievement. "
        "The system accepts a language parameter ('english', 'urdu', or 'bilingual') and "
        "generates both English written digests and Urdu MP3 audio briefings through a two-step "
        "pipeline: an OpenAI-powered translation call followed by Google Text-to-Speech synthesis "
        "using the Urdu language code. This feature directly addresses the needs of Pakistan's "
        "bilingual user base as specified in the SRS."
    )
    add_body(doc,
        "The implementation of content safety guardrails at both input and output boundaries "
        "represents a qualitative improvement over a naive API pass-through architecture. The "
        "GuardrailChecker module enforces a deny-list of prompt injection and harmful content "
        "patterns, returning standardised HTTP 400 responses when violations are detected "
        "rather than forwarding unsafe content to the OpenAI API. Session memory was added to "
        "the AgentRunner, enabling multi-turn conversational continuity within a single browser "
        "session without requiring a persistent database — a lightweight but impactful "
        "improvement to user experience."
    )
    add_body(doc,
        "The frontend achieved a high degree of UI polish through the use of Framer Motion "
        "animations, parallax scrolling effects, a custom cursor follower, and a dark-themed "
        "design language. The live news section implements a BBC-style editorial grid layout "
        "with category filters, an animated live indicator, and a last-updated timestamp. The "
        "Ultimate News Panel and Multi-Agent News Panel represent novel frontend components "
        "not present in comparable open-source news aggregation projects, providing users with "
        "direct access to the system's most powerful capabilities without requiring API "
        "knowledge."
    )

    # ── 7.1  Critical Review ──────────────────────────────────────────────────
    add_heading_custom(doc, "7.1", "Critical Review", 16)

    add_body(doc,
        "A candid review of the AI News Desk reveals several areas where the delivered "
        "implementation diverges from the original SRS specification. The most significant "
        "departure is the agent execution model: while the SRS mandates use of the OpenAI "
        "Agents SDK for agent orchestration, the implementation uses a custom AgentRunner "
        "wrapping the OpenAI Chat Completions API directly. Although the openai-agents package "
        "is listed as a dependency, the sophisticated tool-calling and handoff mechanisms "
        "native to the Agents SDK are not utilised. This design decision was pragmatic — the "
        "Chat Completions API provides greater control over conversation history and session "
        "management — but represents a compliance gap relative to the SRS."
    )
    add_body(doc,
        "The live news feature presents a second critical limitation. The SRS implies delivery "
        "of verified, timestamped news articles from identified sources; the current "
        "implementation generates all news content through LLM inference rather than fetching "
        "live data from RSS feeds, NewsAPI, or a wire service. While the Live News Agent's "
        "system prompt instructs the model to provide real headlines and source URLs, the "
        "outputs are AI-synthesised approximations rather than factually verified live data. "
        "This distinction is important for any deployment context where accuracy is critical."
    )
    add_body(doc,
        "The 'bilingual' mode satisfies REQ-SF5-2 and REQ-SF5-3 at a functional level, but "
        "relies on an intermediate GPT model call to translate English text into Urdu prior to "
        "TTS synthesis. This introduces an additional latency cost and OpenAI API consumption "
        "not accounted for in the original SRS design. The translation quality depends on the "
        "chosen helper model (default: gpt-4o-mini) and may not meet professional journalistic "
        "Urdu standards for all content types, as acknowledged in the SRS business rules."
    )
    add_body(doc,
        "Performance targets stated in the SRS — specifically, sub-8-second single-agent "
        "response times and Lighthouse scores above 80 — were not formally measured or "
        "validated during this project phase. Response times vary significantly based on "
        "OpenAI API load, network conditions, and prompt complexity. The system does not "
        "implement request caching, rate-limit back-off strategies, or fallback to cached "
        "responses, meaning that API quota exhaustion results in complete feature unavailability "
        "rather than degraded-but-functional service."
    )
    add_body(doc,
        "The session management implementation stores conversation history exclusively in "
        "server-side memory, meaning that session context is lost on backend restart. The "
        "SRS notes that the architecture should support migration to a relational database; "
        "however, no migration path or schema versioning mechanism has been implemented. "
        "File-based generated assets (PDF, MP3, PNG) are not associated with authenticated "
        "user accounts, creating a potential privacy concern in a multi-user deployment context."
    )

    # ── 7.2  Lessons Learnt ────────────────────────────────────────────────────
    add_heading_custom(doc, "7.2", "Lessons Learnt", 16)

    add_body(doc,
        "The development of the AI News Desk yielded significant practical insights across "
        "backend architecture, frontend engineering, AI integration, and deployment "
        "configuration. These lessons are documented here to inform future development "
        "iterations and similar projects."
    )
    add_body(doc,
        "From a backend perspective, the most instructive lesson was the importance of "
        "separating concerns clearly between the agent execution layer and the API routing "
        "layer. Early prototypes mixed agent logic directly into FastAPI endpoint handlers, "
        "which made testing and modification difficult. Refactoring into dedicated Agent "
        "subclasses with a shared AgentRunner significantly improved modularity. The "
        "decision to implement use_memory flags on each endpoint — enabling conversational "
        "memory only for single-agent chat routes and disabling it for batch/bulk routes — "
        "prevented subtle context pollution bugs that would otherwise have been difficult to "
        "diagnose."
    )
    add_body(doc,
        "Working with the OpenAI API at scale revealed the critical importance of graceful "
        "quota handling. The project encountered 429 rate-limit and insufficient_quota errors "
        "during testing, which required systematic error classification and appropriate HTTP "
        "status code mapping (402 for quota, 401 for invalid key, 429 for rate limit, 503 "
        "for network failures). This experience reinforced the lesson that external API "
        "dependency management must be a first-class architectural concern, not an afterthought."
    )
    add_body(doc,
        "On the frontend side, the most valuable lesson was the importance of designing for "
        "async state early. The Live News Section, which combines auto-refresh timers, category "
        "filter state, and parsed model output, initially suffered from stale closure bugs in "
        "React's useEffect hooks. Refactoring the fetchNews function with useCallback and "
        "including all dependencies in the effect dependency array resolved these issues. "
        "This reinforced the principle that complex async UI state should be modelled "
        "explicitly rather than relying on implicit closure capture."
    )
    add_body(doc,
        "The deployment phase surfaced important lessons about CORS configuration in "
        "cloud environments. Vercel generates preview deployment URLs with dynamically "
        "assigned subdomains, which are not known at backend deployment time. The initial "
        "approach of hardcoding a single FRONTEND_URL proved insufficient; the team resolved "
        "this by introducing a CORS_EXTRA_ORIGINS environment variable and a CORS_ALLOW_ALL "
        "opt-in flag for preview deployments. This experience highlighted that CORS policy "
        "should be designed with deployment topology in mind from the beginning of the project."
    )
    add_body(doc,
        "Finally, the bilingual TTS feature demonstrated the value of progressive enhancement "
        "in feature design. By wrapping the gTTS call in a try-except block and returning a "
        "voice_error field rather than raising an HTTP 500 error, the system allows the text "
        "result to be returned even when audio generation fails. This pattern — isolating "
        "auxiliary feature failures from core feature delivery — proved to be a robust "
        "design principle applicable across many parts of the system."
    )

    # ── 7.3  Future Enhancements ───────────────────────────────────────────────
    add_heading_custom(doc, "7.3", "Future Enhancements / Recommendations", 16)

    add_body(doc,
        "The following enhancements are recommended to evolve the AI News Desk from a "
        "demonstration-grade system into a production-ready news intelligence platform. "
        "Recommendations are presented in descending order of impact relative to implementation cost."
    )
    add_body(doc,
        "1. Live News Data Pipeline. The highest-impact improvement would be the integration "
        "of real-time news data sources to replace LLM-synthesised news. Recommended "
        "integrations include NewsAPI.org (free tier: 100 requests/day), RSS feed parsing "
        "via the feedparser library for sources such as TechCrunch, MIT Technology Review, "
        "and The Guardian, and optional integration with a commercial news API such as "
        "GDELT or Aylien. This would transform the live news feature from an AI approximation "
        "into a factually verifiable news feed."
    )
    add_body(doc,
        "2. WebSocket-based Real-Time Updates. The current 60-second HTTP polling cycle "
        "creates unnecessary server load and introduces up to 60 seconds of news latency. "
        "Replacing the polling mechanism with a WebSocket connection would enable push-based "
        "live updates from the backend to all connected clients simultaneously, reducing "
        "latency to near-zero and server load by eliminating redundant polling requests."
    )
    add_body(doc,
        "3. User Authentication and Personalisation. Implementing JWT-based user "
        "authentication would enable per-user session persistence (replacing in-memory "
        "session storage), saved news preferences, personal news category subscriptions, "
        "and access to a history of previously generated reports. The SRS architecture "
        "was designed to support this addition without structural changes to the existing "
        "codebase."
    )
    add_body(doc,
        "4. Persistent Storage and Asset Management. Migrating from the current file-system "
        "and in-memory storage to a cloud object store (AWS S3, Cloudflare R2, or similar) "
        "for generated files and a PostgreSQL database for session and agent run records "
        "would provide durability, scalability, and a foundation for analytics. The ERD "
        "presented in Section 4.3 provides a ready-to-implement schema for this migration."
    )
    add_body(doc,
        "5. Full OpenAI Agents SDK Migration. Migrating the Runner to use the official "
        "OpenAI Agents SDK's tool-calling and agent handoff mechanisms would unlock "
        "advanced orchestration capabilities, including multi-step reasoning, tool "
        "invocation (e.g., live web search tools), and structured output parsing. This "
        "would align the implementation with the SRS requirement and enable significantly "
        "richer agent behaviours."
    )
    add_body(doc,
        "6. Caching and Rate-Limit Resilience. Implementing a Redis-based response cache "
        "for common queries would reduce OpenAI API consumption and improve response times "
        "for frequently requested topics. An exponential back-off retry strategy with "
        "configurable maximum retries would improve resilience under API quota pressure, "
        "while a circuit-breaker pattern would prevent cascading failures during extended "
        "API outages."
    )

    # ══════════════════════════════════════════════════════════════════════════
    #  APPENDICES
    # ══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)

    p = doc.add_paragraph()
    run = p.add_run("Appendices")
    set_font(run, "Calibri", 24, bold=True, color=(0x1e, 0x3a, 0x5f))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)

    # ── Appendix A: User Manual ───────────────────────────────────────────────
    add_page_break(doc)
    add_appendix_main_heading(doc, "A", "User Manual")

    add_body(doc,
        "This appendix provides a comprehensive guide for end users of the AI News Desk "
        "platform. It describes the system's user-facing features, explains how to navigate "
        "the web interface, and provides step-by-step instructions for querying AI agents, "
        "using the live news panel, generating PDF and audio reports, and interpreting agent "
        "outputs. The manual assumes no prior technical knowledge and is written for the "
        "general news consumer user class described in SRS Section 2.2.3. Users who "
        "encounter errors are directed to the troubleshooting subsection, which covers "
        "the most common failure modes and their resolutions.",
        font="Times New Roman"
    )

    add_heading_custom(doc, "A.1", "Getting Started", 16)
    add_body(doc,
        "The AI News Desk is accessible at the frontend URL provided by your system "
        "administrator (e.g., https://ai-news-desk.vercel.app). No account registration or "
        "software installation is required. The application runs entirely in your web browser "
        "and supports Chrome 90+, Firefox 88+, Safari 14+, and Edge 90+. Ensure your device "
        "has a stable internet connection before use, as all AI agent features require live "
        "communication with the backend API."
    )

    add_heading_custom(doc, "A.1.1", "Navigating the Homepage", 14)
    add_body(doc,
        "Upon loading the homepage, you will see the Hero Section at the top, featuring "
        "the platform title and two call-to-action buttons: Deploy Agents (takes you to the "
        "full agents directory) and View Demo (opens the agents listing page). Below the "
        "hero, the Multi-Agent News Panel allows you to select multiple agents simultaneously "
        "and run a combined query. The Live News Section appears further down the page, "
        "displaying auto-refreshing AI news with category filter buttons. The Ultimate News "
        "Panel at the bottom of the page provides access to PDF, audio, and graph generation."
    )

    add_heading_custom(doc, "A.1.1.1", "Using the Live News Feature", 12)
    add_body(doc,
        "The live news panel updates automatically every 60 seconds. To filter news by "
        "topic, click one of the category buttons: All, AI/Tech, Crypto, Politics, Health, "
        "Pakistan, Sports, or World. The green animated dot next to the LIVE label indicates "
        "that auto-refresh is active. To pause auto-refresh, uncheck the 'Auto-refresh every "
        "60 seconds' checkbox at the bottom of the panel. Click the Refresh button at any "
        "time to manually fetch the latest news."
    )

    add_heading_custom(doc, "A.1.2", "Using AI Agents", 14)
    add_body(doc,
        "Navigate to the Agents page via the navigation bar or the 'Deploy Agents' button. "
        "You will see a grid of eleven agent cards. Click any agent card to open its dedicated "
        "chat interface. Type your query in the input field at the bottom of the chat and press "
        "Enter or click the Send button. The agent will process your request and return a "
        "formatted response. Each chat session maintains context, so you can ask follow-up "
        "questions without restating the background."
    )

    add_heading_custom(doc, "A.1.2.1", "Generating PDF and Audio Reports", 12)
    add_body(doc,
        "On the homepage, scroll to the Ultimate News Panel section. Enter a news query, "
        "select your desired language (English, Urdu, or Bilingual), and check the features "
        "you want: Daily Collection, Summaries, Trend Graph, PDF Report, and/or Voice Summary. "
        "Click 'Run ultimate agent' and wait for the response (this may take up to 30 seconds). "
        "When complete, download buttons will appear for any generated files. Click 'Download "
        "PDF', 'Download MP3', or 'Download Graph' to save the respective files to your device."
    )

    add_heading_custom(doc, "A.1.3", "Troubleshooting", 14)
    add_body(doc,
        "If you see the message 'OpenAI API quota exceeded', the backend's API credits have "
        "been exhausted. Contact your system administrator to add OpenAI credits. If agent "
        "responses appear empty or generic, verify that the backend server is running by "
        "visiting the /health endpoint. If live news fails to load, check your internet "
        "connection and click the manual Refresh button. For CORS-related errors in the "
        "browser console, ensure the NEXT_PUBLIC_API_URL environment variable on Vercel "
        "matches the deployed backend URL."
    )

    # ── Appendix B: Administrator Manual ─────────────────────────────────────
    add_page_break(doc)
    add_appendix_main_heading(doc, "B", "Administrator Manual")

    add_body(doc,
        "This appendix provides technical guidance for system administrators responsible "
        "for deploying, configuring, and maintaining the AI News Desk platform. It covers "
        "the complete local development setup procedure, cloud deployment instructions for "
        "Vercel and Railway, environment variable configuration, backend dependency "
        "management, CORS configuration for multi-domain deployments, and operational "
        "monitoring guidance. Administrators are assumed to have working knowledge of Python "
        "3.11+, Node.js 18+, Git, and basic cloud platform administration. All commands "
        "are written for a Windows PowerShell environment unless otherwise noted.",
        font="Times New Roman"
    )

    add_heading_custom(doc, "B.1", "Local Development Setup", 16)
    add_body(doc,
        "Clone the repository and navigate to the project root. The project consists of "
        "two independently runnable subsystems: a Python FastAPI backend (backend/) and a "
        "Next.js frontend (frontend/). Both must be running simultaneously for full "
        "functionality. Ensure Python 3.11+ and Node.js 18+ are installed before proceeding."
    )

    add_heading_custom(doc, "B.1.1", "Backend Setup", 14)
    add_body(doc,
        "Install the uv package manager (pip install uv). Navigate to the backend directory "
        "and run: uv pip install -e . This installs all Python dependencies including FastAPI, "
        "Uvicorn, OpenAI SDK, ReportLab, gTTS, and Matplotlib. Create a .env file in the "
        "backend directory with the following variables: OPENAI_API_KEY=sk-your-key-here, "
        "FRONTEND_URL=http://localhost:3000, and ENVIRONMENT=development. Start the backend "
        "with: uv run uvicorn main:app --reload --host 127.0.0.1 --port 8000. The API will "
        "be accessible at http://127.0.0.1:8000 and interactive documentation at /docs."
    )

    add_heading_custom(doc, "B.1.1.1", "Environment Variables Reference", 12)
    add_body(doc,
        "OPENAI_API_KEY (required): Your OpenAI API key from https://platform.openai.com. "
        "FRONTEND_URL (required for production): The full URL of the deployed frontend, "
        "e.g., https://your-app.vercel.app. ENVIRONMENT: Set to 'production' to activate "
        "production CORS rules. CORS_EXTRA_ORIGINS: Comma-separated list of additional "
        "allowed origins for preview deployments. CORS_ALLOW_ALL: Set to '1' to allow all "
        "origins (use only for development or preview). "
        "OPENAI_BILINGUAL_HELPER_MODEL: OpenAI model for Urdu translation (default: gpt-4o-mini)."
    )

    add_heading_custom(doc, "B.1.2", "Frontend Setup", 14)
    add_body(doc,
        "Navigate to the frontend directory and run: npm install. Create a .env.local file "
        "with: NEXT_PUBLIC_API_URL=http://127.0.0.1:8000. Start the development server with: "
        "npm run dev. The frontend will be accessible at http://localhost:3000. If port 3000 "
        "is occupied, Next.js will automatically use the next available port (3001, etc.). "
        "To build a production bundle, run: npm run build, followed by: npm start."
    )

    add_heading_custom(doc, "B.1.2.1", "Production Deployment on Vercel and Railway", 12)
    add_body(doc,
        "For Railway (backend): Create a new project, connect your GitHub repository, set "
        "the root directory to backend/, and configure the start command as: uvicorn main:app "
        "--host 0.0.0.0 --port $PORT. Add all required environment variables in the Railway "
        "dashboard. Railway will provide a public HTTPS URL. For Vercel (frontend): Import "
        "the repository, set the root directory to frontend/, and add the environment "
        "variable NEXT_PUBLIC_API_URL pointing to your Railway backend URL. Redeploy after "
        "adding environment variables. Update the backend's FRONTEND_URL to match your "
        "Vercel deployment URL and redeploy the backend."
    )

    add_heading_custom(doc, "B.1.3", "Monitoring and Maintenance", 14)
    add_body(doc,
        "Health check: GET /health returns {status: healthy} when the backend is operational. "
        "Generated files (PDF, MP3, PNG) accumulate in the reports/, audio/, and graphs/ "
        "directories on the backend server. The startup cleanup hook deletes files older than "
        "24 hours automatically on each server restart. For manual cleanup, run: "
        "cleanup_old_generated_files(max_age_hours=24) from utils/helpers.py. Monitor OpenAI "
        "API quota usage at https://platform.openai.com/account/usage. Backend logs are "
        "available in the Railway dashboard under the Deployments tab."
    )

    # ── Appendix C: Information / Promotional Material ───────────────────────
    add_page_break(doc)
    add_appendix_main_heading(doc, "C", "Information / Promotional Material")

    add_body(doc,
        "This appendix contains the promotional and informational materials developed to "
        "present the AI News Desk platform to prospective users, academic evaluators, and "
        "industry stakeholders at the project showcase and public demonstration events. "
        "The materials include a product brochure, a concise event flyer, a free-standing "
        "display standee, and a large-format banner — each designed to communicate the "
        "platform's core value proposition and visual identity. All materials follow the "
        "project's established colour palette (deep navy #1e3a5f, cyan #06b6d4, and purple "
        "#7c3aed) and use Calibri as the primary typeface to maintain brand consistency "
        "across print and digital formats.",
        font="Times New Roman"
    )

    add_heading_custom(doc, "C.1", "Brochure", 16)
    add_body(doc,
        "The AI News Desk product brochure is a bi-fold A4 document designed for distribution "
        "at academic showcases and technology exhibitions. The front panel features the platform "
        "logo, the tagline 'News at the Speed of Artificial Intelligence', and a QR code linking "
        "to the live deployment. The inner panels describe the six AI agents (SEO, YouTube, "
        "Forbes, Web Search, Research, Breaking News) with brief capability summaries, the "
        "live news feature with its eight category filters, and the Ultimate Agent's PDF/audio/"
        "graph output modalities. The back panel lists the technology stack (FastAPI, Next.js, "
        "OpenAI, ReportLab, gTTS, Matplotlib), the deployment infrastructure (Vercel + Railway), "
        "and contact information."
    )

    add_heading_custom(doc, "C.2", "Flyer", 16)
    add_body(doc,
        "The event flyer is a single A5 sheet designed for rapid audience engagement at the "
        "project demonstration booth. It features a bold headline ('AI News Desk — 11 Agents. "
        "One Platform.'), three key benefit statements ('Real-time AI news', 'PDF + Audio "
        "reports', 'English & Urdu support'), the live deployment URL, and a QR code. The "
        "flyer uses a dark background with gradient text to reflect the application's UI "
        "aesthetic and is designed for both digital distribution (via email and social media) "
        "and physical printing."
    )

    add_heading_custom(doc, "C.3", "Standee", 16)
    add_body(doc,
        "The pull-up standee is an 85 cm x 200 cm vertical display designed for the project "
        "showcase floor. The upper section features the AI News Desk logo and tagline at large "
        "scale. The middle section presents a simplified architecture diagram showing the "
        "three-tier deployment (Browser → Vercel → Railway → OpenAI) and a feature highlights "
        "list. The lower section includes team member names, supervisor details, and the "
        "institution logo. The standee is designed in portrait orientation with bleed marks "
        "for professional printing."
    )

    add_heading_custom(doc, "C.4", "Banner", 16)
    add_body(doc,
        "The horizontal banner is a 200 cm x 75 cm landscape display for use above the "
        "demonstration table. It carries the platform name in large bold type on the left, "
        "a row of six agent icons with labels in the centre, and the institution name and "
        "project year on the right. The banner uses the same navy-to-purple gradient "
        "background as the application's hero section, creating visual continuity between "
        "the physical display and the on-screen demonstration."
    )

    add_heading_custom(doc, "C.5", "Design Assets and Brand Guidelines", 16)
    add_body(doc,
        "All promotional materials were created using the following brand specifications: "
        "primary colour #1e3a5f (deep navy), accent colour #06b6d4 (cyan), secondary accent "
        "#7c3aed (purple), background #030014 (near-black). Primary typeface: Calibri (headings) "
        "and Inter (body text). Logo: stylised AI circuit motif with the text 'AI News Desk' "
        "in Outfit semibold. All print materials were prepared at 300 DPI minimum resolution. "
        "Digital assets are available in SVG and PNG formats."
    )

    add_heading_custom(doc, "C.1.1", "Digital Distribution Strategy", 14)
    add_body(doc,
        "Digital versions of the brochure and flyer were distributed via the project's "
        "GitHub repository README, the institution's project showcase webpage, and the "
        "platform's About page accessible from the frontend navigation. The QR codes in "
        "all materials link to the live Vercel deployment URL, ensuring that audiences can "
        "access the platform immediately from any mobile device at the showcase event."
    )

    add_heading_custom(doc, "C.1.1.1", "QR Code Specifications", 12)
    add_body(doc,
        "All QR codes were generated at error correction level H (30% damage tolerance) "
        "to ensure reliable scanning under varying lighting conditions at the showcase venue. "
        "The QR codes encode the full HTTPS URL of the Vercel deployment and were tested "
        "across Android and iOS native camera apps as well as three third-party QR scanner "
        "applications to confirm consistent decoding."
    )

    # ── Appendix D: API Reference ─────────────────────────────────────────────
    add_page_break(doc)
    add_appendix_main_heading(doc, "D", "API Reference")

    add_body(doc,
        "This appendix provides a concise reference for all API endpoints exposed by the "
        "AI News Desk backend. It is intended for developers integrating the backend into "
        "third-party systems or building alternative frontend clients. Each endpoint entry "
        "describes the HTTP method, path, request body schema, and response body schema. "
        "Full interactive documentation is available at the /docs endpoint of any running "
        "backend instance (powered by FastAPI's automatic Swagger UI generation). "
        "All endpoints accept and return JSON-encoded bodies and require the "
        "Content-Type: application/json header on POST requests.",
        font="Times New Roman"
    )

    add_heading_custom(doc, "D.1", "Core Agent Endpoints", 16)
    add_body(doc,
        "POST /api/agent: Runs a single agent. Request body: {query: string, agent_type: "
        "string, session_id?: string}. Response: {result: string, session_id: string, "
        "agent_type: string}. Valid agent_type values: seo, youtube, forbes, web_search, "
        "news_research, breaking_news_alert, news_summarizer, multi_agent_newsroom, "
        "ultimate_ai_news, live_news. Session memory is maintained across calls sharing "
        "the same session_id.\n\n"
        "POST /api/news: Runs multiple agents concurrently. Request body: {query: string, "
        "agents: string[], session_id?: string}. Response: {results: {[agent]: string}, "
        "session_id: string}."
    )

    add_heading_custom(doc, "D.1.1", "Specialised Endpoints", 14)
    add_body(doc,
        "POST /api/live-news: Returns live AI news. Request body: {categories?: string[], "
        "session_id?: string}. Valid categories: all, ai_tech, crypto, politics, health, "
        "pakistan, sports, world.\n\n"
        "POST /api/ultimate-news: Runs the Ultimate Agent with optional file generation. "
        "Request body: {query: string, features?: string[], language?: string}. Features: "
        "daily_collection, summaries, trends, trend_graph, pdf_report, voice_summary. "
        "Language: english, urdu, bilingual. Response includes pdf_path, voice_path, "
        "graph_path fields when respective features are enabled."
    )

    add_heading_custom(doc, "D.1.1.1", "File Download Endpoints", 12)
    add_body(doc,
        "GET /api/download-pdf/{filename}: Serves a generated PDF file. Returns "
        "Content-Type: application/pdf. Returns HTTP 404 if file does not exist, HTTP 400 "
        "if filename is invalid.\n\n"
        "GET /api/download-audio/{filename}: Serves a generated MP3 file. Returns "
        "Content-Type: audio/mpeg.\n\n"
        "GET /api/download-graph/{filename}: Serves a generated PNG trend graph. Returns "
        "Content-Type: image/png.\n\n"
        "GET /health: Returns {status: healthy} with HTTP 200 when backend is operational."
    )

    # ══════════════════════════════════════════════════════════════════════════
    #  REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_ref_heading(doc, "Reference and Bibliography")

    refs = [
        ("1",  'OpenAI. (2024). OpenAI Agents SDK Documentation. Available at: https://platform.openai.com/docs/agents. [Accessed: May 2026].'),
        ("2",  'S. Ramirez. (2024). FastAPI Official Documentation, Version 0.111. Tiangolo. Available at: https://fastapi.tiangolo.com. [Accessed: May 2026].'),
        ("3",  'Vercel Inc. (2024). Next.js Deployment Guide. Available at: https://vercel.com/docs. [Accessed: May 2026].'),
        ("4",  'ReportLab Group. (2024). ReportLab User\'s Guide, Version 4.0. Available at: https://www.reportlab.com/docs/reportlab-userguide.pdf. [Accessed: May 2026].'),
        ("5",  'L. Quast. (2024). gTTS (Google Text-to-Speech) Python Library Documentation, Version 2.5. Available at: https://gtts.readthedocs.io. [Accessed: May 2026].'),
        ("6",  'IEEE. (1998). IEEE Std 830-1998: Recommended Practice for Software Requirements Specifications. Institute of Electrical and Electronics Engineers, New York, USA.'),
        ("7",  'Matplotlib Development Team. (2024). Matplotlib 3.8 Documentation. Available at: https://matplotlib.org/stable/index.html. [Accessed: May 2026].'),
        ("8",  'Pydantic. (2024). Pydantic v2 Documentation. Available at: https://docs.pydantic.dev. [Accessed: May 2026].'),
        ("9",  'Railway Corp. (2024). Railway Deployment Documentation. Available at: https://docs.railway.app. [Accessed: May 2026].'),
        ("10", 'M. Lutz. (2019). Learning Python, 5th Edition. O\'Reilly Media, Sebastopol, USA. ISBN: 978-1449355739.'),
        ("11", 'A. Banks, E. Porcello. (2020). Learning React: Modern Patterns for Developing React Apps. O\'Reilly Media, Sebastopol, USA. ISBN: 978-1492051725.'),
        ("12", 'OpenAI. (2023). GPT-4 Technical Report. arXiv preprint arXiv:2303.08774. Available at: https://arxiv.org/abs/2303.08774.'),
        ("13", 'T. Brown et al. (2020). "Language Models are Few-Shot Learners", Advances in Neural Information Processing Systems (NeurIPS), Volume 33, pp. 1877-1901.'),
        ("14", 'Astral. (2024). uv: An Extremely Fast Python Package Installer and Resolver. Available at: https://docs.astral.sh/uv. [Accessed: May 2026].'),
        ("15", 'Google. (2024). Google Cloud Text-to-Speech API Documentation. Available at: https://cloud.google.com/text-to-speech/docs. [Accessed: May 2026].'),
    ]

    for tag, text in refs:
        add_reference(doc, tag, text)

    # ══════════════════════════════════════════════════════════════════════════
    #  INDEX
    # ══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_ref_heading(doc, "Index")

    index_entries = {
        "A": [
            "Activity Diagram  ·  4.7",
            "Agent Manager  ·  4.9",
            "AgentRunner  ·  4.4, 5.1.1",
            "asyncio.gather()  ·  5.1.1, 7.1",
            "Audio generation (MP3)  ·  2.4.4, 5.2",
        ],
        "B": [
            "Backend (FastAPI)  ·  2.2.1, 4.1, 5.3",
            "Bilingual support  ·  2.4.5, 7.1, 7.3",
            "Breaking News Agent  ·  2.4.1, 3.1",
        ],
        "C": [
            "Calibri (font)  ·  throughout",
            "ChatBot component  ·  4.9",
            "Component Diagram  ·  4.9",
            "Content safety / Guardrails  ·  2.5.2, 4.6",
            "CORS  ·  2.5.3, 5.3, 7.2",
        ],
        "D": [
            "Data Flow Diagram  ·  4.11",
            "Deployment Diagram  ·  4.10",
            "Daily News Collector Agent  ·  2.4.1",
        ],
        "E": [
            "Environment variables  ·  2.2.5, 5.3, B.1.1.1",
        ],
        "F": [
            "FastAPI  ·  2.2.1, 2.3.3, 4.1",
            "Forbes Agent  ·  2.4.1",
            "Frontend (Next.js)  ·  2.2.1, 4.1, 5.3",
        ],
        "G": [
            "gTTS (Google Text-to-Speech)  ·  2.4.4, 5.2",
            "GuardrailChecker  ·  2.5.2, 4.4, 4.6",
        ],
        "L": [
            "Live News  ·  2.4.2, 4.7, 7.1",
            "LiveNewsSection component  ·  4.9",
        ],
        "M": [
            "Matplotlib  ·  2.4.3, 5.2",
            "Multi-Agent Newsroom  ·  2.4.1",
            "Multi-Agent News Panel  ·  4.9",
        ],
        "N": [
            "Next.js  ·  2.2.1, 2.3.3, 5.2",
        ],
        "O": [
            "OpenAI Agents SDK  ·  2.3.3, 5.2, 7.1",
            "OpenAI API  ·  2.2.7, 4.10, 4.11",
        ],
        "P": [
            "PDF Report Generation  ·  2.4.3, 5.1.2",
            "Python / uv  ·  2.2.4, 5.4",
        ],
        "R": [
            "Railway (cloud)  ·  2.2.4, 4.10",
            "ReportLab  ·  2.4.3, 5.2",
            "Research Agent  ·  2.4.1",
        ],
        "S": [
            "SEO Agent  ·  2.4.1",
            "Session management  ·  2.4.1, 4.8, 5.1.1",
            "State Transition Diagram  ·  4.8",
        ],
        "T": [
            "Tailwind CSS  ·  2.2.5, 5.2",
            "TypeScript  ·  2.2.5, 5.2",
        ],
        "U": [
            "Ultimate AI News Agent  ·  2.4.3, 5.1.2, 7.1",
            "Urdu / bilingual  ·  2.4.5, 7.1",
            "Use Case Analysis  ·  Ch. 3",
        ],
        "V": [
            "Vercel (CDN)  ·  2.2.4, 4.10, B.1.2.1",
            "Voice Summary  ·  2.4.4, 5.1.2",
        ],
        "W": [
            "Web Search Agent  ·  2.4.1",
            "WebSocket (future)  ·  7.3",
        ],
        "Y": [
            "YouTube Agent  ·  2.4.1",
        ],
    }

    for letter, entries in index_entries.items():
        p = doc.add_paragraph()
        r = p.add_run(f"[{letter}]")
        set_font(r, "Calibri", 12, bold=True, color=(0x1e, 0x3a, 0x5f))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        for entry in entries:
            ep = doc.add_paragraph(style="List Bullet")
            er = ep.add_run(entry)
            set_font(er, "Calibri", 11)
            ep.paragraph_format.space_after  = Pt(1)
            ep.paragraph_format.line_spacing = Pt(15)

    # save
    out = r"c:\Users\DELL\Desktop\ai-news\Chapter7_Appendices.docx"
    doc.save(out)
    print(f"Done! Saved -> {out}")


if __name__ == "__main__":
    build()

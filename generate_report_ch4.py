"""
Generate Chapter 4 (4.7–4.11) Word document with inline diagrams.
Run from project root: uv run python generate_report_ch4.py
"""

import io
import os

import matplotlib.patches as mpatches
import matplotlib.patches as mpl_patches
import matplotlib.pyplot as plt
import matplotlib.patheffects as pe
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import matplotlib.patches as patches

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ────────────────────────────────────────────────────────────────────

def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def box(ax, x, y, w, h, label, color="#1e3a5f", text_color="white",
        fontsize=8.5, radius=0.04, border="#0d2137"):
    rect = FancyBboxPatch(
        (x - w / 2, y - h / 2), w, h,
        boxstyle=f"round,pad=0.02,rounding_size={radius}",
        linewidth=1.2, edgecolor=border, facecolor=color, zorder=3,
    )
    ax.add_patch(rect)
    ax.text(x, y, label, ha="center", va="center",
            color=text_color, fontsize=fontsize,
            fontweight="bold", zorder=4, wrap=True,
            multialignment="center")


def diamond(ax, x, y, w, h, label, color="#e0a820", text_color="white", fontsize=8):
    dx, dy = w / 2, h / 2
    poly = plt.Polygon(
        [[x, y + dy], [x + dx, y], [x, y - dy], [x - dx, y]],
        closed=True, linewidth=1.2, edgecolor="#8b6500", facecolor=color, zorder=3,
    )
    ax.add_patch(poly)
    ax.text(x, y, label, ha="center", va="center",
            color=text_color, fontsize=fontsize, fontweight="bold", zorder=4,
            multialignment="center")


def arr(ax, x1, y1, x2, y2, label="", color="#555", lw=1.4):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=color,
                                lw=lw, mutation_scale=12), zorder=2)
    if label:
        mx, my = (x1 + x2) / 2 + 0.02, (y1 + y2) / 2
        ax.text(mx, my, label, fontsize=7, color="#333", ha="left", va="center")


def hbar(ax, x1, x2, y, color="#555", lw=3):
    ax.plot([x1, x2], [y, y], color=color, linewidth=lw, zorder=3)


def oval(ax, x, y, rx, ry, label, color="#2e7d32", text_color="white", fontsize=9):
    ell = mpatches.Ellipse((x, y), rx, ry,
                            linewidth=1.2, edgecolor="#1b5e20",
                            facecolor=color, zorder=3)
    ax.add_patch(ell)
    ax.text(x, y, label, ha="center", va="center",
            color=text_color, fontsize=fontsize,
            fontweight="bold", zorder=4)


# ── set Word styles ─────────────────────────────────────────────────────────────

def set_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    return p


def body_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = "Calibri"
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(18)
    return p


def insert_fig(doc, stream, width_inches=6.0, caption=""):
    doc.add_picture(stream, width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

# ── 4.7  Activity Diagram ──────────────────────────────────────────────────────

def make_activity_diagram():
    fig, axes = plt.subplots(1, 2, figsize=(13, 11))
    fig.patch.set_facecolor("#f5f7fa")

    # ── LEFT: Ultimate Agent Workflow ────────────────────────────────────────
    ax = axes[0]
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_facecolor("#f5f7fa")
    ax.set_title("Ultimate AI News Agent Workflow", fontsize=10, fontweight="bold",
                 color="#1e3a5f", pad=8)

    # nodes  (x, y)
    N = {
        "start":  (0.50, 0.95),
        "recv":   (0.50, 0.87),
        "guard1": (0.50, 0.79),
        "ok1":    (0.50, 0.71),   # decision: input ok?
        "collect":(0.30, 0.61),
        "trend":  (0.70, 0.61),
        "sync1":  (0.50, 0.52),
        "summ":   (0.50, 0.44),
        "ok2":    (0.50, 0.36),   # decision: pdf?
        "pdf":    (0.28, 0.27),
        "ok3":    (0.72, 0.36),   # decision: voice?
        "voice":  (0.72, 0.27),
        "assemble":(0.50, 0.19),
        "guard2": (0.50, 0.12),
        "resp":   (0.50, 0.05),
    }

    # start oval
    oval(ax, *N["start"], 0.18, 0.05, "START", color="#263238")
    box(ax, *N["recv"],  0.52, 0.06, "Receive\n/api/ultimate-news", color="#1565c0")
    box(ax, *N["guard1"],0.52, 0.06, "Validate Input\n(Guardrail Check)", color="#6a1b9a")
    diamond(ax, *N["ok1"], 0.40, 0.07, "Input\nSafe?")
    # parallel fork bar
    hbar(ax, 0.16, 0.84, 0.605, color="#1e3a5f")
    box(ax, *N["collect"],0.38, 0.07, "Daily News\nCollector Agent", color="#00695c")
    box(ax, *N["trend"],  0.38, 0.07, "Generate\nTrend Graph\n(Matplotlib)", color="#00695c")
    # sync bar
    hbar(ax, 0.16, 0.84, 0.475, color="#1e3a5f")
    box(ax, *N["summ"],   0.52, 0.06, "Summarize News\n(TLDR)", color="#1565c0")

    # decisions
    diamond(ax, *N["ok2"], 0.38, 0.065, "PDF\nEnabled?")
    box(ax,   *N["pdf"],   0.36, 0.065, "Generate PDF\n(ReportLab)", color="#bf360c")
    diamond(ax, *N["ok3"], 0.38, 0.065, "Voice\nEnabled?")
    box(ax,   *N["voice"], 0.36, 0.065, "Generate Audio\n(gTTS)", color="#bf360c")

    box(ax, *N["assemble"],0.52, 0.06, "Assemble JSON\nResponse", color="#1565c0")
    box(ax, *N["guard2"],  0.52, 0.06, "Output\nGuardrail Check", color="#6a1b9a")
    oval(ax, *N["resp"],   0.18, 0.05, "API RESPONSE", color="#263238")

    # arrows
    for a, b in [("start","recv"),("recv","guard1"),("guard1","ok1")]:
        arr(ax, N[a][0], N[a][1]-0.028, N[b][0], N[b][1]+0.028)

    # ok1 → fork bar (via "Yes")
    arr(ax, N["ok1"][0], N["ok1"][1]-0.04, 0.50, 0.61, label="Yes")
    ax.text(0.22, 0.605, "parallel\nfork", fontsize=6.5, color="#1e3a5f",
            ha="center", va="center")
    # ok1 → 400 (No)
    ax.annotate("", xy=(0.92, N["ok1"][1]),
                xytext=(N["ok1"][0]+0.20, N["ok1"][1]),
                arrowprops=dict(arrowstyle="-|>", color="#c62828", lw=1.3))
    ax.text(0.84, N["ok1"][1]+0.012, "No → 400", fontsize=6.5, color="#c62828")

    # fork→collect and fork→trend
    arr(ax, 0.30, 0.605, N["collect"][0], N["collect"][1]+0.035)
    arr(ax, 0.70, 0.605, N["trend"][0],   N["trend"][1]+0.035)

    # collect,trend → sync bar
    arr(ax, N["collect"][0], N["collect"][1]-0.035, 0.30, 0.475)
    arr(ax, N["trend"][0],   N["trend"][1]-0.035,   0.70, 0.475)

    ax.text(0.22, 0.475, "join", fontsize=6.5, color="#1e3a5f", ha="center")
    arr(ax, 0.50, 0.475, N["summ"][0], N["summ"][1]+0.03)
    arr(ax, N["summ"][0], N["summ"][1]-0.03, N["ok2"][0], N["ok2"][1]+0.03)

    # ok2 yes → pdf
    arr(ax, N["ok2"][0]-0.19, N["ok2"][1], N["pdf"][0]+0.18, N["pdf"][1], label="Yes")
    # ok2 no → ok3
    arr(ax, N["ok2"][0]+0.19, N["ok2"][1], N["ok3"][0]-0.19, N["ok3"][1], label="No")
    # ok3 yes → voice
    arr(ax, N["ok3"][0], N["ok3"][1]-0.035, N["voice"][0], N["voice"][1]+0.035, label="Yes")

    # pdf & voice → assemble
    arr(ax, N["pdf"][0],   N["pdf"][1]-0.035,   N["assemble"][0]-0.05, N["assemble"][1]+0.03)
    arr(ax, N["voice"][0], N["voice"][1]-0.035, N["assemble"][0]+0.05, N["assemble"][1]+0.03)
    arr(ax, N["ok3"][0],   N["ok3"][1]-0.035,   N["assemble"][0]+0.10, N["assemble"][1]+0.03)

    arr(ax, N["assemble"][0], N["assemble"][1]-0.03, N["guard2"][0], N["guard2"][1]+0.03)
    arr(ax, N["guard2"][0],   N["guard2"][1]-0.03,   N["resp"][0],   N["resp"][1]+0.025)

    # ── RIGHT: Single-Agent Workflow ──────────────────────────────────────────
    ax2 = axes[1]
    ax2.set_xlim(0, 1)
    ax2.set_ylim(0, 1)
    ax2.axis("off")
    ax2.set_facecolor("#f5f7fa")
    ax2.set_title("Single-Agent / Multi-Agent Query Workflow", fontsize=10,
                  fontweight="bold", color="#1e3a5f", pad=8)

    steps = [
        (0.50, 0.92, "START",                    "oval",    "#263238"),
        (0.50, 0.82, "Receive Request",           "box",     "#1565c0"),
        (0.50, 0.72, "Validate Input\n(Guardrail)","box",    "#6a1b9a"),
        (0.50, 0.62, "Input Safe?",               "diamond", "#e0a820"),
        (0.50, 0.51, "Run Selected Agents\n(parallel swimlanes for multi-agent)", "box", "#00695c"),
        (0.50, 0.41, "Collect All Results\n(sync bar)", "box", "#1565c0"),
        (0.50, 0.31, "Validate Output\n(Guardrail)", "box",  "#6a1b9a"),
        (0.50, 0.21, "Assemble Response",         "box",     "#1565c0"),
        (0.50, 0.11, "Return API Response",       "oval",    "#263238"),
    ]

    for i, (x, y, lbl, kind, col) in enumerate(steps):
        if kind == "oval":
            oval(ax2, x, y, 0.32, 0.055, lbl, color=col)
        elif kind == "box":
            box(ax2, x, y, 0.62, 0.065, lbl, color=col)
        else:
            diamond(ax2, x, y, 0.44, 0.075, lbl)

        if i > 0:
            py = steps[i-1][1]
            arr(ax2, x, py - 0.034, x, y + 0.034)

    # "No → 400" on decision
    ax2.annotate("", xy=(0.92, 0.62), xytext=(0.72, 0.62),
                 arrowprops=dict(arrowstyle="-|>", color="#c62828", lw=1.3))
    ax2.text(0.78, 0.635, "No → 400", fontsize=6.5, color="#c62828")
    ax2.text(0.53, 0.565, "Yes", fontsize=6.5, color="#333")

    # swimlane note
    ax2.text(0.50, 0.475, "← Agent A  |  Agent B  |  Agent C →",
             fontsize=6, color="#555", ha="center", style="italic")

    fig.tight_layout(pad=2.5)
    return fig_to_stream(fig)


# ── 4.8  State Transition Diagram ─────────────────────────────────────────────

def make_state_diagram():
    fig, ax = plt.subplots(figsize=(11, 7))
    fig.patch.set_facecolor("#f5f7fa")
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_facecolor("#f5f7fa")
    ax.set_title("4.8  Agent Session State Transition Diagram",
                 fontsize=11, fontweight="bold", color="#1e3a5f", pad=10)

    states = {
        "Initialized": (2.0, 5.5, "#1565c0"),
        "Active":       (5.5, 5.5, "#2e7d32"),
        "File Generation":(5.5, 3.5, "#e65100"),
        "Blocked":      (9.0, 5.5, "#b71c1c"),
        "Expired":      (5.5, 1.5, "#37474f"),
    }

    for name, (x, y, col) in states.items():
        box(ax, x, y, 2.2, 0.9, name, color=col, fontsize=9.5)

    def tr(x1, y1, x2, y2, label="", color="#444", rad=0.0, lw=1.5):
        style = f"arc3,rad={rad}" if rad else "arc3,rad=0"
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color,
                                   connectionstyle=style, lw=lw))
        mx, my = (x1+x2)/2, (y1+y2)/2
        if label:
            ax.text(mx, my+0.12, label, fontsize=7.2, ha="center",
                    color=color, fontweight="bold",
                    bbox=dict(fc="white", ec="none", pad=1))

    # start pseudo-state
    ax.add_patch(plt.Circle((0.6, 5.5), 0.18, color="#263238", zorder=4))
    tr(0.78, 5.5, 2.0-1.1, 5.5, "First API call\n(UUID generated)", "#263238")

    tr(3.1, 5.5,  4.4, 5.5, "Valid query received", "#2e7d32")
    tr(5.5+1.1, 5.5, 9.0-1.1, 5.5, "Guardrail violation", "#b71c1c")
    tr(9.0-1.1, 5.5, 5.5+1.1, 5.5, "Session reset", "#2e7d32", rad=0.35)
    tr(5.5, 5.5-0.45, 5.5, 3.5+0.45, "File generation\nrequested", "#e65100")
    tr(5.5+0.5, 3.5+0.45, 5.5+0.5, 5.5-0.45, "File generation\ncomplete", "#2e7d32", rad=0.4)
    tr(5.5, 3.5-0.45, 5.5, 1.5+0.45, "30-min timeout\n(no activity)", "#37474f")
    tr(5.5, 5.5-0.45, 5.5, 1.5+0.45, "Timeout", "#37474f", rad=0.45)
    tr(2.0, 5.5-0.45, 5.5-1.1, 1.5+0.2, "Timeout", "#37474f", rad=0.2)

    # self-loop Active → Active
    loop = mpatches.FancyArrowPatch(
        (5.5+1.1, 5.8), (5.5+1.1, 6.2),
        connectionstyle="arc3,rad=-0.8",
        arrowstyle="-|>", color="#2e7d32", linewidth=1.4, zorder=3)
    ax.add_patch(loop)
    ax.text(7.5, 6.05, "Next query\n(stays Active)", fontsize=7, color="#2e7d32",
            ha="center", fontweight="bold")

    # New session from Expired
    ax.add_patch(plt.Circle((5.5, 0.55), 0.18, color="#263238", zorder=4))
    ax.add_patch(plt.Circle((5.5, 0.55), 0.22, color="none",
                             ec="#263238", lw=2, zorder=3))
    tr(5.5, 1.5-0.45, 5.5, 0.73, "Next call →\nnew session", "#263238")

    ax.text(0.5, 0.3, "● = start pseudo-state   ◎ = end (Expired → new session)",
            fontsize=7.5, color="#555", style="italic")

    return fig_to_stream(fig)


# ── 4.9  Component Diagram ─────────────────────────────────────────────────────

def make_component_diagram():
    fig, ax = plt.subplots(figsize=(13, 8))
    fig.patch.set_facecolor("#f5f7fa")
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor("#f5f7fa")
    ax.set_title("4.9  Component Diagram – AI News Desk",
                 fontsize=11, fontweight="bold", color="#1e3a5f", pad=10)

    # subsystem backgrounds
    be = FancyBboxPatch((0.3, 0.4), 5.6, 7.0,
                        boxstyle="round,pad=0.1", lw=1.5,
                        edgecolor="#1565c0", facecolor="#e3f2fd", zorder=0)
    fe = FancyBboxPatch((6.3, 0.4), 4.0, 7.0,
                        boxstyle="round,pad=0.1", lw=1.5,
                        edgecolor="#2e7d32", facecolor="#e8f5e9", zorder=0)
    ex = FancyBboxPatch((10.7, 0.4), 2.1, 7.0,
                        boxstyle="round,pad=0.1", lw=1.5,
                        edgecolor="#e65100", facecolor="#fff3e0", zorder=0)
    for p in [be, fe, ex]:
        ax.add_patch(p)

    ax.text(3.1, 7.2, "BACKEND", fontsize=10, fontweight="bold",
            color="#1565c0", ha="center")
    ax.text(8.3, 7.2, "FRONTEND", fontsize=10, fontweight="bold",
            color="#2e7d32", ha="center")
    ax.text(11.75, 7.2, "EXTERNAL", fontsize=10, fontweight="bold",
            color="#e65100", ha="center")

    # backend components
    bcomps = {
        "API Router\n(main.py)":          (3.1, 6.2, "#1565c0"),
        "Agent Manager\n(agents/)":       (1.6, 4.8, "#00695c"),
        "GuardrailChecker\n(guardrails.py)":(4.6, 4.8, "#6a1b9a"),
        "AgentRunner\n(agent_runner.py)": (1.6, 3.4, "#00695c"),
        "Utility Services\n(utils/helpers.py)":(3.1, 2.0, "#e65100"),
        "File Server\n(download endpoints)":(3.1, 0.9, "#37474f"),
        "SessionStore\n(session_store.py)":(4.6, 3.4, "#37474f"),
    }
    for lbl, (x, y, col) in bcomps.items():
        box(ax, x, y, 2.4, 0.65, lbl, color=col, fontsize=7.5)

    # frontend components
    fcomps = {
        "HeroSection":          (8.3, 6.2, "#1976d2"),
        "AgentsGrid +\nAgentCard": (8.3, 5.2, "#1976d2"),
        "ChatBot":              (8.3, 4.2, "#1976d2"),
        "LiveNewsSection":      (8.3, 3.2, "#1976d2"),
        "MultiAgentNewsPanel":  (8.3, 2.2, "#1976d2"),
        "UltimateNewsPanel":    (8.3, 1.2, "#1976d2"),
        "API Client\n(lib/api.ts)": (8.3, 0.2, "#43a047"),
    }
    for lbl, (x, y, col) in fcomps.items():
        box(ax, x, y, 3.2, 0.6, lbl, color=col, fontsize=7.5)

    # external
    ecomps = {
        "OpenAI\nAgents SDK":  (11.75, 5.5, "#e65100"),
        "ReportLab\n(PDF)":    (11.75, 4.3, "#bf360c"),
        "gTTS\n(MP3)":         (11.75, 3.1, "#bf360c"),
        "Matplotlib\n(Graph)": (11.75, 1.9, "#bf360c"),
        "Google TTS\nService": (11.75, 0.7, "#bf360c"),
    }
    for lbl, (x, y, col) in ecomps.items():
        box(ax, x, y, 1.8, 0.6, lbl, color=col, fontsize=7.5)

    def dep(ax, x1, y1, x2, y2, col="#555", style="arc3,rad=0", lw=1.4, dash=False):
        ls = (0, (4, 3)) if dash else "solid"
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=col,
                                   connectionstyle=style, lw=lw,
                                   linestyle=ls))

    # API Router ↔ Agent Manager
    dep(ax, 3.1, 6.2-0.33, 1.6, 4.8+0.33, col="#1565c0")
    dep(ax, 3.1, 6.2-0.33, 4.6, 4.8+0.33, col="#6a1b9a")
    dep(ax, 3.1, 6.2-0.33, 3.1, 2.0+0.33, col="#e65100")
    dep(ax, 3.1, 6.2-0.33, 3.1, 0.9+0.33, col="#37474f")
    dep(ax, 1.6, 4.8-0.33, 1.6, 3.4+0.33, col="#00695c")
    dep(ax, 4.6, 3.4-0.33, 3.1, 2.0+0.33, col="#6a1b9a")

    # Agent Manager → OpenAI SDK
    dep(ax, 1.6+1.2, 4.8, 11.75-0.9, 5.5, col="#e65100", dash=True)
    # Utility → ReportLab, gTTS, Matplotlib
    dep(ax, 3.1+1.2, 2.0, 11.75-0.9, 4.3, col="#bf360c", dash=True)
    dep(ax, 3.1+1.2, 2.0, 11.75-0.9, 3.1, col="#bf360c", dash=True)
    dep(ax, 3.1+1.2, 2.0, 11.75-0.9, 1.9, col="#bf360c", dash=True)
    dep(ax, 3.1+1.2, 2.0, 11.75-0.9, 0.7, col="#bf360c", dash=True, style="arc3,rad=-0.2")

    # API Client → API Router (primary inter-system link)
    dep(ax, 8.3-1.6, 0.2, 3.1+1.2, 6.2-0.33, col="#e91e63", lw=2.0)
    ax.text(5.7, 3.2, "HTTP/JSON\n(HTTPS)", fontsize=7, color="#e91e63",
            ha="center", fontweight="bold",
            bbox=dict(fc="white", ec="#e91e63", pad=2, lw=0.8))

    # legend
    legend_items = [
        mpatches.Patch(color="#1565c0", label="Backend component"),
        mpatches.Patch(color="#2e7d32", label="Frontend component"),
        mpatches.Patch(color="#e65100", label="External service"),
        mpatches.Patch(color="#e91e63", label="Primary inter-system dependency"),
    ]
    ax.legend(handles=legend_items, loc="lower right",
              fontsize=7.5, framealpha=0.9)

    return fig_to_stream(fig)


# ── 4.10  Deployment Diagram ───────────────────────────────────────────────────

def make_deployment_diagram():
    fig, ax = plt.subplots(figsize=(13, 7.5))
    fig.patch.set_facecolor("#f5f7fa")
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7.5)
    ax.axis("off")
    ax.set_facecolor("#f5f7fa")
    ax.set_title("4.10  Deployment Diagram – AI News Desk",
                 fontsize=11, fontweight="bold", color="#1e3a5f", pad=10)

    nodes = [
        (1.6,  3.75, 2.8, 5.8, "#e3f2fd", "#1565c0", "Node 1: Vercel CDN",
         ["Next.js Build Output", "Static Assets (SSR/CDN)", "ENV: NEXT_PUBLIC_API_URL",
          "Protocol: HTTPS (TLS 1.2+)"]),
        (5.15, 3.75, 2.8, 5.8, "#e8f5e9", "#2e7d32", "Node 2: Railway Cloud",
         ["FastAPI App (Uvicorn ASGI)", "ENV: OPENAI_API_KEY, FRONTEND_URL",
          "reports/ audio/ graphs/ (file store)", "Protocol: HTTPS"]),
        (8.7,  3.75, 2.8, 5.8, "#fff3e0", "#e65100", "Node 3: OpenAI API Cloud",
         ["GPT-4o Inference", "Chat Completions API", "Agents SDK Endpoint",
          "Protocol: HTTPS"]),
        (1.6,  0.4,  2.8, 2.8, "#fce4ec", "#c62828", "Client (User Browser)",
         ["Chrome / Firefox / Safari", "ES2020+ JavaScript",
          "Viewport: 375px – 1920px"]),
    ]

    for cx, cy, w, h, fc, ec, title, items in nodes:
        r = FancyBboxPatch((cx - w/2, cy - h/2), w, h,
                           boxstyle="round,pad=0.08", lw=2,
                           edgecolor=ec, facecolor=fc, zorder=1)
        ax.add_patch(r)
        ax.text(cx, cy + h/2 - 0.28, title, ha="center", va="top",
                fontsize=8.5, fontweight="bold", color=ec)
        for k, item in enumerate(items):
            ax.text(cx, cy + h/2 - 0.62 - k*0.38, f"• {item}",
                    ha="center", va="top", fontsize=7, color="#333")

    def node_arr(x1, y1, x2, y2, label="", col="#555", rad=0.0, lw=1.8):
        style = f"arc3,rad={rad}"
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="<->", color=col,
                                   connectionstyle=style, lw=lw))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx, my+0.15, label, fontsize=7.5, ha="center",
                    color=col, fontweight="bold",
                    bbox=dict(fc="white", ec=col, pad=2, lw=0.8))

    # Client ↔ Vercel
    node_arr(1.6, 0.4+1.4, 1.6, 3.75-2.9, "HTTPS\n(browser)", "#1565c0", lw=2)
    # Vercel ↔ Railway
    node_arr(1.6+1.4, 3.75, 5.15-1.4, 3.75, "REST / JSON\nHTTPS", "#2e7d32", lw=2)
    # Railway ↔ OpenAI
    node_arr(5.15+1.4, 3.75, 8.7-1.4, 3.75, "OpenAI SDK\nHTTPS", "#e65100", lw=2)

    # note: client does NOT reach OpenAI directly
    ax.annotate("", xy=(8.7-1.4, 0.9), xytext=(1.6+1.4, 0.9),
                arrowprops=dict(arrowstyle="-", color="#aaa",
                                connectionstyle="arc3,rad=0", lw=1,
                                linestyle=(0, (4, 3))))
    ax.text(5.15, 0.72, "Client never contacts OpenAI directly (credentials server-side)",
            ha="center", fontsize=6.8, color="#888", style="italic")

    ax.text(11.5, 5.5, "Security Note:\nAPI keys stored\nonly on Node 2\n(env vars)",
            fontsize=8, color="#b71c1c", ha="center",
            bbox=dict(fc="#ffebee", ec="#b71c1c", pad=5, lw=1))

    return fig_to_stream(fig)


# ── 4.11  Data Flow Diagram ────────────────────────────────────────────────────

def make_dfd():
    fig, axes = plt.subplots(1, 2, figsize=(14, 7))
    fig.patch.set_facecolor("#f5f7fa")

    # ── Level 0 (Context) ────────────────────────────────────────────────────
    ax = axes[0]
    ax.set_xlim(0, 7)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_facecolor("#f5f7fa")
    ax.set_title("Level 0: Context DFD", fontsize=10,
                 fontweight="bold", color="#1e3a5f", pad=8)

    # external entities (rectangles)
    box(ax, 1.0, 3.5, 1.6, 0.8, "User", color="#37474f", fontsize=10)
    box(ax, 6.0, 3.5, 1.6, 0.8, "OpenAI API", color="#e65100", fontsize=10)
    # central process (circle)
    circ = plt.Circle((3.5, 3.5), 1.0, color="#1565c0",
                      ec="#0d2137", lw=1.5, zorder=3)
    ax.add_patch(circ)
    ax.text(3.5, 3.5, "AI News\nDesk\nSystem", ha="center", va="center",
            color="white", fontsize=9, fontweight="bold", zorder=4)

    def flow(ax, x1, y1, x2, y2, label, col="#333", lw=1.4, fs=7):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=col, lw=lw))
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx, my+0.2, label, ha="center", fontsize=fs,
                color=col, fontweight="bold",
                bbox=dict(fc="white", ec="none", pad=1))

    # User → System
    flow(ax, 1.8, 3.8, 2.5, 3.6, "News Queries\n& Agent Selections", "#1565c0")
    # System → User (above)
    flow(ax, 2.5, 3.4, 1.8, 3.2, "News Results\nGenerated Reports\nLive Updates", "#2e7d32")
    # System → OpenAI
    flow(ax, 4.5, 3.6, 5.2, 3.6, "Agent Prompts\n& Requests", "#e65100")
    # OpenAI → System
    flow(ax, 5.2, 3.4, 4.5, 3.4, "AI-Generated\nText Responses", "#e65100")

    # ── Level 1 DFD ──────────────────────────────────────────────────────────
    ax2 = axes[1]
    ax2.set_xlim(0, 7)
    ax2.set_ylim(0, 7)
    ax2.axis("off")
    ax2.set_facecolor("#f5f7fa")
    ax2.set_title("Level 1: DFD – Internal Processes", fontsize=10,
                  fontweight="bold", color="#1e3a5f", pad=8)

    # external entities
    box(ax2, 0.7, 5.8, 1.2, 0.55, "User",       color="#37474f", fontsize=8)
    box(ax2, 0.7, 1.0, 1.2, 0.55, "User",       color="#37474f", fontsize=8)
    box(ax2, 6.3, 4.0, 1.2, 0.55, "OpenAI API", color="#e65100", fontsize=8)

    # data store
    ax2.add_patch(FancyBboxPatch((2.5, 2.2), 2.0, 0.55,
                                 boxstyle="round,pad=0.05", lw=1.2,
                                 edgecolor="#555", facecolor="#f9fbe7"))
    ax2.text(3.5, 2.47, "File Store\n(reports/ audio/ graphs/)",
             ha="center", va="center", fontsize=7.5, color="#333")

    # processes
    procs = {
        "1.0 Receive &\nValidate Request": (3.5, 5.8, "#1565c0"),
        "2.0 Execute\nAgent Pipeline":     (3.5, 4.5, "#00695c"),
        "3.0 Generate\nAuxiliary Files":   (3.5, 3.3, "#e65100"),
        "4.0 Assemble &\nReturn Response": (3.5, 1.0, "#1565c0"),
        "5.0 Monitor\nLive News":          (1.4, 4.0, "#6a1b9a"),
    }
    for lbl, (x, y, col) in procs.items():
        circ2 = plt.Circle((x, y), 0.62, color=col, ec="#0d2137",
                           lw=1.2, zorder=3)
        ax2.add_patch(circ2)
        ax2.text(x, y, lbl, ha="center", va="center", color="white",
                 fontsize=6.5, fontweight="bold", zorder=4,
                 multialignment="center")

    # flows
    flow(ax2, 1.3, 5.8, 2.88, 5.8,  "Query",        "#1565c0", fs=6.5)
    flow(ax2, 2.88, 4.9, 2.88, 5.18, "Validated",    "#2e7d32", fs=6.5)
    flow(ax2, 3.5,  5.18,3.5,  5.12, "",             "#555")
    flow(ax2, 3.5,  4.88, 3.5, 4.62, "Agent tasks",  "#00695c", fs=6.5)
    flow(ax2, 3.5+0.62, 4.5, 5.7, 4.0, "Inference",  "#e65100", fs=6.5)
    flow(ax2, 5.7, 4.0, 3.5+0.62, 4.5, "AI Text",    "#e65100", fs=6.5)
    flow(ax2, 3.5, 3.88, 3.5, 3.62,  "Agent Output", "#e65100", fs=6.5)
    flow(ax2, 3.5, 2.97, 3.5, 2.75,  "File refs",    "#555",    fs=6.5)
    flow(ax2, 3.5, 2.2,  3.5, 1.62,  "Metadata",     "#555",    fs=6.5)
    flow(ax2, 2.88, 1.0, 1.3, 1.0,   "Response",     "#2e7d32", fs=6.5)
    # 5.0 → 2.0
    flow(ax2, 1.4, 4.62, 2.88, 4.5,  "Auto-poll", "#6a1b9a", fs=6.5)

    fig.tight_layout(pad=2.5)
    return fig_to_stream(fig)


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── 4.7 ─────────────────────────────────────────────────────────────────
    set_heading(doc, "4.7   Activity Diagram", level=2)
    body_para(doc,
        "The activity diagram for the Ultimate AI News Agent workflow describes the sequence "
        "of activities performed when a user requests a comprehensive news digest. The workflow "
        "begins when the user submits a request to the /api/ultimate-news endpoint. The system "
        "first runs the Daily News Collector activity, which queries the OpenAI agent for news "
        "across all monitored topics. Concurrently, if the trends feature is requested, the "
        "system executes the Generate Trend Graph activity using Matplotlib. Following news "
        "collection, the Summarize News activity generates TLDR summaries of the collected "
        "articles. If the PDF feature is enabled, the Generate PDF Report activity creates a "
        "ReportLab-formatted document. If the voice feature is enabled, the Generate Audio "
        "Summary activity calls gTTS to create an MP3 file. Once all selected activities are "
        "complete, the Assemble Response activity collects all results, filenames, and metadata "
        "into a structured JSON object. The Guardrail Check activity validates the assembled "
        "response before it is returned to the frontend. The workflow concludes with the API "
        "response being sent to the client."
    )
    body_para(doc,
        "For the simpler single-agent query workflow, the activity diagram is linear: "
        "Receive Request → Validate Input (Guardrail) → Run Selected Agent(s) → Collect "
        "Results → Validate Output (Guardrail) → Return Response. Concurrent agent execution "
        "is represented as parallel swimlanes for multi-agent queries, with a synchronization "
        "bar collecting all agent results before response assembly."
    )
    print("Generating Activity Diagram...")
    insert_fig(doc, make_activity_diagram(), 6.2,
               "Figure 4.7 – Activity Diagram: Ultimate Agent Workflow (left) and "
               "Single/Multi-Agent Query Workflow (right)")
    doc.add_paragraph()

    # ── 4.8 ─────────────────────────────────────────────────────────────────
    set_heading(doc, "4.8   State Transition Diagram", level=2)
    body_para(doc,
        "The state transition diagram for the agent session lifecycle describes the states "
        "that a user session passes through during interaction with the AI News Desk. A session "
        "begins in the Initialized state upon the user's first API call, when a UUID session ID "
        "is generated. Upon receiving the first valid query, the session transitions to the "
        "Active state, in which agent runs are executed and results are returned. If the user "
        "makes another query within the session, the state remains Active. If the user requests "
        "file generation (PDF, audio, or graph), the session enters the File Generation state, "
        "a transient state that returns to Active upon completion. If the backend detects no "
        "activity for a configurable timeout period (default: 30 minutes), the session "
        "transitions to the Expired state, and session context is cleared from memory. Sessions "
        "in the Expired state cannot be resumed; a new session is created on the next API call. "
        "If a guardrail violation is detected, the session enters the Blocked state, in which "
        "subsequent queries are rejected until the session is reset."
    )
    print("Generating State Transition Diagram...")
    insert_fig(doc, make_state_diagram(), 6.2,
               "Figure 4.8 – State Transition Diagram: Agent Session Lifecycle")
    doc.add_paragraph()

    # ── 4.9 ─────────────────────────────────────────────────────────────────
    set_heading(doc, "4.9   Component Diagram", level=2)
    body_para(doc,
        "The component diagram for the AI News Desk describes the modular components of the "
        "system and their dependencies. The backend system consists of the following components: "
        "the API Router component (main.py), which receives all HTTP requests and delegates to "
        "the appropriate handler; the Agent Manager component (agents/), which instantiates and "
        "executes agent classes; the Guardrail component (embedded in the Runner), which performs "
        "input/output safety checks; the Utility Services component (utils/helpers.py), which "
        "provides PDF, audio, and graph generation functions; and the File Server component, "
        "which handles file download endpoints. The frontend system consists of the Hero "
        "Component (landing page search interface), the AgentsGrid and AgentCard components "
        "(agent selection cards), the ChatBot component (per-agent chat interface), the "
        "LiveNewsSection Component (auto-refreshing news panel), the MultiAgentNewsPanel, the "
        "UltimateNewsPanel, and the API Client module (lib/api.ts), which manages all HTTP calls "
        "to the backend. The API Client depends on the backend API Router, creating the primary "
        "inter-subsystem dependency. The Agent Manager depends on the OpenAI Agents SDK "
        "(external). Utility Services depend on ReportLab, gTTS, and Matplotlib (external libraries)."
    )
    print("Generating Component Diagram...")
    insert_fig(doc, make_component_diagram(), 6.2,
               "Figure 4.9 – Component Diagram: Backend, Frontend, and External Services")
    doc.add_paragraph()

    # ── 4.10 ────────────────────────────────────────────────────────────────
    set_heading(doc, "4.10   Deployment Diagram", level=2)
    body_para(doc,
        "The deployment diagram for the AI News Desk describes the physical distribution of "
        "system components across hardware and cloud nodes. The system is deployed across three "
        "primary nodes. Node 1 is the Vercel CDN, a globally distributed content delivery "
        "network that hosts the compiled Next.js frontend application. Vercel serves static "
        "assets and server-side rendered pages to end users' browsers via HTTPS. The deployed "
        "artifact on this node is the Next.js build output, configured with the "
        "NEXT_PUBLIC_API_URL environment variable pointing to the backend URL. Node 2 is the "
        "Railway Cloud Container, a containerized Linux environment that hosts the FastAPI "
        "backend application. The backend server runs as a Uvicorn ASGI process on a "
        "dynamically assigned port. This node also hosts the reports/, audio/, and graphs/ "
        "directories for generated file storage. The deployed artifact is the Python backend "
        "application with all dependencies installed via uv. Node 3 is the OpenAI API Cloud, "
        "the external service that provides AI inference capabilities. The backend communicates "
        "with this node via HTTPS using the OpenAI Python SDK. End users' browsers communicate "
        "only with Node 1; the backend (Node 2) communicates with Node 3. This topology ensures "
        "that API credentials remain server-side and the client never directly accesses the AI "
        "inference service."
    )
    print("Generating Deployment Diagram...")
    insert_fig(doc, make_deployment_diagram(), 6.2,
               "Figure 4.10 – Deployment Diagram: Three-Node Cloud Architecture")
    doc.add_paragraph()

    # ── 4.11 ────────────────────────────────────────────────────────────────
    set_heading(doc, "4.11   Data Flow Diagram", level=2)
    body_para(doc,
        "The Level 0 (Context) Data Flow Diagram for the AI News Desk shows the system as a "
        "single process box (the AI News Desk System) with two external entities: the User and "
        "the OpenAI API. Data flows from the User to the system in the form of News Queries and "
        "Agent Selections. Data flows from the system to the User in the form of News Results, "
        "Generated Reports (PDF, MP3, PNG), and Live News Updates. Data flows from the system "
        "to the OpenAI API in the form of Agent Prompts and Inference Requests. Data flows from "
        "the OpenAI API to the system in the form of AI-Generated Text Responses."
    )
    body_para(doc,
        "The Level 1 Data Flow Diagram decomposes the AI News Desk System into its primary "
        "processes. Process 1.0 (Receive and Validate Request) accepts the user query, validates "
        "it against guardrails, and routes it to the appropriate process. Process 2.0 (Execute "
        "Agent Pipeline) sends inference requests to the OpenAI API and collects agent outputs. "
        "Process 3.0 (Generate Auxiliary Files) receives agent outputs and triggers PDF, audio, "
        "or graph generation as required, storing results in the File Store data store. Process "
        "4.0 (Assemble and Return Response) collects all results, retrieves file metadata from "
        "the File Store, and returns the complete response to the user. The Live News Update "
        "flow is handled by a parallel Process 5.0 (Monitor Live News), which periodically "
        "triggers Process 2.0 for the live news agent and pushes updates to the user's browser "
        "via the polling mechanism."
    )
    print("Generating Data Flow Diagram...")
    insert_fig(doc, make_dfd(), 6.2,
               "Figure 4.11 – Data Flow Diagram: Level 0 Context (left) and "
               "Level 1 Internal Processes (right)")

    out = r"c:\Users\DELL\Desktop\ai-news\Chapter4_Diagrams.docx"
    doc.save(out)
    print(f"\nDone! Document saved -> {out}")


if __name__ == "__main__":
    build_doc()
